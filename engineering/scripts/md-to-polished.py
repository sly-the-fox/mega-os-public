#!/usr/bin/env python3
"""
md-to-polished.py — Convert Markdown documents to polished DOCX/PDF output.

Usage:
    python md-to-polished.py <input.md> [--format docx|pdf|both] [--output-dir DIR]

Requires: python-docx, markdown-it-py
Optional: libreoffice (for PDF conversion)
"""

import argparse
import os
import re
import subprocess
import sys
from datetime import date
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT
except ImportError:
    print("Error: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)

try:
    from markdown_it import MarkdownIt
except ImportError:
    print("Error: markdown-it-py not installed. Run: pip install markdown-it-py")
    sys.exit(1)


REPO_ROOT = Path(__file__).resolve().parent.parent.parent
DEFAULT_OUTPUT_DIR = REPO_ROOT / "deliverables"


def clean_text(text: str) -> str:
    """Remove markdown artifacts and clean em dashes."""
    # Remove bold/italic markers
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'__(.+?)__', r'\1', text)
    text = re.sub(r'_(.+?)_', r'\1', text)
    # Remove inline code backticks
    text = re.sub(r'`(.+?)`', r'\1', text)
    # Remove links, keep text
    text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
    # Clean em dashes: replace " — " with ", " or " - "
    text = text.replace(' — ', ', ')
    text = text.replace('—', ', ')
    # Clean horizontal rules
    text = text.replace('---', '')
    text = text.replace('***', '')
    text = text.replace('___', '')
    return text.strip()


def extract_bold_runs(text: str):
    """Extract text with bold markers preserved as run instructions.

    Returns list of (text, is_bold) tuples.
    """
    runs = []
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            inner = part[2:-2]
            inner = inner.replace(' — ', ', ').replace('—', ', ')
            runs.append((inner, True))
        else:
            cleaned = part.replace(' — ', ', ').replace('—', ', ')
            if cleaned:
                runs.append((cleaned, False))
    return runs


def parse_markdown(md_text: str):
    """Parse markdown into structured blocks."""
    blocks = []
    lines = md_text.split('\n')
    i = 0
    title = None

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            i += 1
            continue

        # Horizontal rules
        if re.match(r'^-{3,}$|^\*{3,}$|^_{3,}$', stripped):
            i += 1
            continue

        # Headings
        heading_match = re.match(r'^(#{1,6})\s+(.+)', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = clean_text(heading_match.group(2))
            if level == 1 and title is None:
                title = text
            blocks.append({'type': 'heading', 'level': level, 'text': text})
            i += 1
            continue

        # Bullet lists
        bullet_match = re.match(r'^(\s*)[-*+]\s+(.+)', line)
        if bullet_match:
            indent = len(bullet_match.group(1))
            level = indent // 2
            text = bullet_match.group(2)
            blocks.append({'type': 'bullet', 'text': text, 'level': level})
            i += 1
            continue

        # Numbered lists
        num_match = re.match(r'^(\s*)\d+\.\s+(.+)', line)
        if num_match:
            indent = len(num_match.group(1))
            level = indent // 2
            text = num_match.group(2)
            blocks.append({'type': 'numbered', 'text': text, 'level': level})
            i += 1
            continue

        # Block quotes
        quote_match = re.match(r'^>\s*(.*)', stripped)
        if quote_match:
            text = clean_text(quote_match.group(1))
            blocks.append({'type': 'quote', 'text': text})
            i += 1
            continue

        # Regular paragraph — collect consecutive non-empty lines
        para_lines = [stripped]
        i += 1
        while i < len(lines):
            next_stripped = lines[i].strip()
            if not next_stripped:
                break
            if re.match(r'^#{1,6}\s+', next_stripped):
                break
            if re.match(r'^[-*+]\s+', next_stripped):
                break
            if re.match(r'^\d+\.\s+', next_stripped):
                break
            if re.match(r'^>\s*', next_stripped):
                break
            if re.match(r'^-{3,}$|^\*{3,}$|^_{3,}$', next_stripped):
                break
            para_lines.append(next_stripped)
            i += 1

        full_text = ' '.join(para_lines)
        blocks.append({'type': 'paragraph', 'text': full_text})

    return title, blocks


def add_formatted_text(paragraph, text: str):
    """Add text to a paragraph, preserving bold markers as formatting."""
    # Remove link syntax first
    text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
    # Remove inline code
    text = re.sub(r'`(.+?)`', r'\1', text)

    runs = extract_bold_runs(text)
    for run_text, is_bold in runs:
        # Clean remaining markdown from non-bold text
        if not is_bold:
            run_text = re.sub(r'\*(.+?)\*', r'\1', run_text)
            run_text = re.sub(r'_(.+?)_', r'\1', run_text)
        run = paragraph.add_run(run_text)
        run.bold = is_bold


def build_docx(title: str, blocks: list, output_path: Path):
    """Build a DOCX document from parsed blocks."""
    doc = Document()

    # Configure default style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x1C, 0x1C, 0x1C)
    paragraph_format = style.paragraph_format
    paragraph_format.space_after = Pt(6)
    paragraph_format.line_spacing = 1.15

    # Configure heading styles
    for i in range(1, 5):
        heading_style = doc.styles[f'Heading {i}']
        heading_style.font.name = 'Calibri'
        heading_style.font.color.rgb = RGBColor(0x1C, 0x1C, 0x1C)
        if i == 1:
            heading_style.font.size = Pt(24)
            heading_style.paragraph_format.space_before = Pt(0)
            heading_style.paragraph_format.space_after = Pt(12)
        elif i == 2:
            heading_style.font.size = Pt(18)
            heading_style.paragraph_format.space_before = Pt(18)
            heading_style.paragraph_format.space_after = Pt(8)
        elif i == 3:
            heading_style.font.size = Pt(14)
            heading_style.paragraph_format.space_before = Pt(12)
            heading_style.paragraph_format.space_after = Pt(6)
        else:
            heading_style.font.size = Pt(12)
            heading_style.paragraph_format.space_before = Pt(10)
            heading_style.paragraph_format.space_after = Pt(4)

    # Title page
    if title:
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.paragraph_format.space_before = Pt(120)
        title_para.paragraph_format.space_after = Pt(24)
        run = title_para.add_run(title)
        run.font.size = Pt(28)
        run.font.name = 'Calibri'
        run.bold = True

        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_para.paragraph_format.space_after = Pt(60)
        run = date_para.add_run(date.today().strftime('%B %d, %Y'))
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        doc.add_page_break()

    # Add content
    for block in blocks:
        btype = block['type']

        if btype == 'heading':
            level = min(block['level'], 4)
            # Skip H1 if it matches title (already on cover page)
            if level == 1 and block['text'] == title:
                continue
            doc.add_heading(block['text'], level=level)

        elif btype == 'paragraph':
            para = doc.add_paragraph()
            add_formatted_text(para, block['text'])

        elif btype == 'bullet':
            para = doc.add_paragraph(style='List Bullet')
            if block.get('level', 0) > 0:
                para.paragraph_format.left_indent = Inches(0.5 * block['level'])
            add_formatted_text(para, block['text'])

        elif btype == 'numbered':
            para = doc.add_paragraph(style='List Number')
            if block.get('level', 0) > 0:
                para.paragraph_format.left_indent = Inches(0.5 * block['level'])
            add_formatted_text(para, block['text'])

        elif btype == 'quote':
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.5)
            run = para.add_run(block['text'])
            run.italic = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Add page numbers in footer
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Page number field
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    fld_char_begin = OxmlElement('w:fldChar')
    fld_char_begin.set(qn('w:fldCharType'), 'begin')
    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = ' PAGE '
    fld_char_end = OxmlElement('w:fldChar')
    fld_char_end.set(qn('w:fldCharType'), 'end')

    run = footer_para.add_run()
    run._element.append(fld_char_begin)
    run2 = footer_para.add_run()
    run2._element.append(instr_text)
    run3 = footer_para.add_run()
    run3._element.append(fld_char_end)

    doc.save(str(output_path))
    return output_path


def convert_to_pdf(docx_path: Path, output_dir: Path) -> Path:
    """Convert DOCX to PDF using libreoffice."""
    try:
        result = subprocess.run(
            ['libreoffice', '--headless', '--convert-to', 'pdf',
             '--outdir', str(output_dir), str(docx_path)],
            capture_output=True, text=True, timeout=60
        )
        if result.returncode != 0:
            print(f"Warning: libreoffice conversion failed: {result.stderr}")
            return None
        pdf_path = output_dir / docx_path.with_suffix('.pdf').name
        return pdf_path if pdf_path.exists() else None
    except FileNotFoundError:
        print("Warning: libreoffice not found. Skipping PDF conversion.")
        return None
    except subprocess.TimeoutExpired:
        print("Warning: libreoffice conversion timed out.")
        return None


def main():
    parser = argparse.ArgumentParser(
        description='Convert Markdown to polished DOCX/PDF'
    )
    parser.add_argument('input', help='Input markdown file path')
    parser.add_argument(
        '--format', choices=['docx', 'pdf', 'both'], default='both',
        help='Output format (default: both)'
    )
    parser.add_argument(
        '--output-dir', type=Path, default=DEFAULT_OUTPUT_DIR,
        help=f'Output directory (default: {DEFAULT_OUTPUT_DIR})'
    )
    args = parser.parse_args()

    input_path = Path(args.input).resolve()
    if not input_path.exists():
        print(f"Error: {input_path} not found")
        sys.exit(1)

    # Read and parse
    md_text = input_path.read_text(encoding='utf-8')
    title, blocks = parse_markdown(md_text)
    if not title:
        title = input_path.stem.replace('-', ' ').title()

    # Prepare output
    args.output_dir.mkdir(parents=True, exist_ok=True)
    date_prefix = date.today().strftime('%Y-%m-%d')
    base_name = f"{date_prefix}-{input_path.stem}"

    results = []

    # Generate DOCX
    if args.format in ('docx', 'both'):
        docx_path = args.output_dir / f"{base_name}.docx"
        build_docx(title, blocks, docx_path)
        size = docx_path.stat().st_size / 1024
        print(f"DOCX: {docx_path} ({size:.1f} KB)")
        results.append(docx_path)

    # Generate PDF
    if args.format in ('pdf', 'both'):
        # Need DOCX first for PDF conversion
        if args.format == 'pdf':
            docx_path = args.output_dir / f"{base_name}.docx"
            build_docx(title, blocks, docx_path)

        pdf_path = convert_to_pdf(docx_path, args.output_dir)
        if pdf_path:
            size = pdf_path.stat().st_size / 1024
            print(f"PDF:  {pdf_path} ({size:.1f} KB)")
            results.append(pdf_path)

        # Clean up temp DOCX if only PDF was requested
        if args.format == 'pdf' and docx_path.exists():
            docx_path.unlink()

    if not results:
        print("No output generated.")
        sys.exit(1)


if __name__ == '__main__':
    main()
